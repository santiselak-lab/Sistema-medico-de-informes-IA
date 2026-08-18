from datetime import datetime, timedelta
import io
import json
import re
import unicodedata
from docx import Document
from openai import OpenAI
from supabase import create_client
import streamlit as st

st.set_page_config(
    page_title="Sistema Médico - Expediente Clínico", layout="wide"
)

# Conexión y credenciales
try:
    groq_key = st.secrets["GROQ_API_KEY"]
    supabase_url = st.secrets["SUPABASE_URL"]
    supabase_key = st.secrets["SUPABASE_KEY"]

    client = OpenAI(
        base_url="https://api.groq.com/openai/v1", api_key=groq_key
    )
    supabase = create_client(supabase_url, supabase_key)
except Exception:
    st.error("⚠️ Revisa las credenciales en 'Secrets' de Streamlit Cloud.")
    st.stop()


def obtener_modelo_chat_activo():
    """Selección dinámica del modelo en Groq para evitar errores 404."""
    modelos_preferidos = [
        "llama-3.3-70b-versatile",
        "llama-3.1-8b-instant",
        "llama3-70b-8192",
        "mixtral-8x7b-32768",
    ]
    try:
        modelos_disponibles = [m.id for m in client.models.list().data]
        for m in modelos_preferidos:
            if m in modelos_disponibles:
                return m
        for m in modelos_disponibles:
            if "whisper" not in m:
                return m
    except Exception:
        pass
    return "llama-3.1-8b-instant"


def formatear_fecha(iso_str):
    """Ajuste de zona horaria UTC a hora local (-6 hrs)"""
    try:
        dt = datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
        dt_local = dt - timedelta(hours=6)
        return dt_local.strftime("%d/%m/%Y %H:%M hrs")
    except Exception:
        return iso_str[:16]


def generar_documento_word(paciente, fecha, transcripcion):
    """Genera documento Word con los datos vigentes de la base de datos"""
    doc = Document()
    doc.add_heading(f"Informe Clínico — {paciente}", level=1)
    doc.add_paragraph(f"Fecha de consulta: {fecha}")
    doc.add_heading("Transcripción Médica:", level=2)
    doc.add_paragraph(transcripcion)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def sanitizar_nombre_archivo(texto):
    """Elimina acentos y caracteres especiales para evitar errores en Supabase Storage"""
    texto_sin_acentos = (
        unicodedata.normalize("NFKD", texto)
        .encode("ASCII", "ignore")
        .decode("utf-8")
    )
    limpio = re.sub(r"[^\w\s-]", "", texto_sin_acentos).strip().replace(" ", "_")
    return limpio if limpio else "Archivo"


st.sidebar.title("🏥 Sistema Médico")
rol = st.sidebar.radio(
    "Selecciona tu Rol:", ["👨‍⚕️ Vista Médico", "👩‍💼 Vista Secretaria"]
)

# =========================================================
# VISTA MÉDICO
# =========================================================
if rol == "👨‍⚕️ Vista Médico":
    st.title("👨‍⚕️ Captura de Audio Médico")

    audio_grabado = st.audio_input("🎙️ Opción A: Grabar dictado en vivo")
    audio_subido = st.file_uploader(
        "📁 Opción B: Subir nota de voz HD (.m4a, .mp3, .wav)",
        type=["m4a", "mp3", "wav"],
    )

    audio_data = audio_grabado if audio_grabado is not None else audio_subido

    if audio_data is not None:
        st.audio(audio_data)

        if st.button(
            "🚀 PROCESAR Y ENVIAR", type="primary", use_container_width=True
        ):
            with st.status(
                "Procesando dictado médico...", expanded=True
            ) as status:
                try:
                    status.write("⏳ Leyendo archivo de audio...")
                    audio_bytes = audio_data.getvalue()
                    buffer_audio = io.BytesIO(audio_bytes)
                    ext = (
                        audio_data.name.split(".")[-1]
                        if hasattr(audio_data, "name")
                        else "wav"
                    )
                    buffer_audio.name = f"dictado.{ext}"

                    status.write("🎙️ Transcribiendo con Whisper...")
                    transcripcion_bruta = client.audio.transcriptions.create(
                        model="whisper-large-v3",
                        file=buffer_audio,
                        language="es",
                        prompt=(
                            "Dictado clínico formal médico. Incluye nombres de"
                            " pacientes, DNI, expedientes, fosa ilíaca, "
                            "juntura, tumoración, diagnóstico y tratamiento."
                        ),
                    ).text

                    modelo_activo = obtener_modelo_chat_activo()
                    status.write("🩺 Optimizando términos médicos con IA...")

                    transcripcion_pulida = transcripcion_bruta
                    try:
                        system_prompt = (
                            "Eres un editor médico en español. Tu ÚNICA función"
                            " es corregir ortografía y términos médicos del"
                            " dictado. NUNCA expliques tus pasos, NUNCA"
                            " analices las instrucciones, NUNCA respondas en"
                            " inglés. Devuelve EXCLUSIVAMENTE el texto"
                            " dictado corregido."
                        )
                        prompt_correccion = f"Dictado a corregir:\n{transcripcion_bruta}"

                        res_corr = client.chat.completions.create(
                            model=modelo_activo,
                            messages=[
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": prompt_correccion},
                            ],
                            temperature=0.0,
                        )
                        texto_salida = res_corr.choices[0].message.content.strip()

                        # Filtro de seguridad por si el LLM incluye análisis en inglés
                        if (
                            "Analyze User Input" not in texto_salida
                            and "**Role:**" not in texto_salida
                        ):
                            transcripcion_pulida = texto_salida
                    except Exception:
                        pass

                    status.write("🧠 Identificando paciente...")
                    nombre_paciente = "Paciente Desconocido"
                    try:
                        prompt_json = f"""
Extrae el NOMBRE COMPLETO del paciente mencionado. Responde ÚNICAMENTE en JSON con la clave "paciente".
Dictado: "{transcripcion_pulida}"
"""
                        res_json = client.chat.completions.create(
                            model=modelo_activo,
                            messages=[
                                {"role": "user", "content": prompt_json}
                            ],
                            response_format={"type": "json_object"},
                            temperature=0.0,
                        )
                        nombre_paciente = (
                            json.loads(res_json.choices[0].message.content)
                            .get("paciente", "Paciente Desconocido")
                            .strip()
                        )
                    except Exception:
                        pass

                    if (
                        nombre_paciente == "Paciente Desconocido"
                        or not nombre_paciente
                    ):
                        match = re.search(
                            r"paciente\s+([A-ZÁÉÍÓÚÑa-záléíóúñ\s]{3,35})(?:,|\s+DNI|\s+del|\s+con|\s+de)",
                            transcripcion_pulida,
                            re.IGNORECASE,
                        )
                        if match:
                            nombre_paciente = match.group(1).strip().title()

                    status.write("☁️ Guardando audio en Supabase...")
                    nombre_limpio = sanitizar_nombre_archivo(nombre_paciente)
                    timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                    file_name = f"{timestamp_str}_{nombre_limpio}.{ext}"

                    supabase.storage.from_("audios").upload(
                        file_name,
                        audio_bytes,
                        file_options={
                            "content-type": f"audio/{ext}",
                            "upsert": "true",
                        },
                    )
                    audio_url = supabase.storage.from_(
                        "audios"
                    ).get_public_url(file_name)

                    status.write("💾 Registrando en base de datos...")
                    supabase.table("informes").insert({
                        "paciente": nombre_paciente,
                        "transcripcion": transcripcion_pulida,
                        "audio_url": audio_url,
                    }).execute()

                    status.update(
                        label="✅ ¡Proceso completado!",
                        state="complete",
                        expanded=False,
                    )
                    st.success(
                        f"¡Consulta guardada para **{nombre_paciente}**!"
                    )

                except Exception as e:
                    status.update(label="❌ Error en el proceso", state="error")
                    st.error(f"Detalle técnico: {e}")

# =========================================================
# VISTA SECRETARIA
# =========================================================
elif rol == "👩‍💼 Vista Secretaria":
    st.title("👩‍💼 Portafolio de Gestión y Expedientes")
    try:
        respuesta = (
            supabase.table("informes")
            .select("*")
            .order("created_at", desc=True)
            .execute()
        )
        datos = respuesta.data

        if not datos:
            st.info("No hay informes registrados aún.")
        else:
            # Agrupar consultas por paciente
            pacientes_dict = {}
            for d in datos:
                nombre = d["paciente"]
                if nombre not in pacientes_dict:
                    pacientes_dict[nombre] = []
                pacientes_dict[nombre].append(d)

            col_pac, col_vis = st.columns(2)
            with col_pac:
                paciente_sel = st.selectbox(
                    "👤 Seleccionar Paciente:", list(pacientes_dict.keys())
                )

            consultas = pacientes_dict[paciente_sel]

            with col_vis:
                opciones_consultas = [
                    f"📅 Consulta: {formatear_fecha(c['created_at'])}"
                    for c in consultas
                ]
                idx_c = st.selectbox(
                    "🕒 Seleccionar Consulta del Historial:",
                    range(len(opciones_consultas)),
                    format_func=lambda x: opciones_consultas[x],
                )

            informe_actual = consultas[idx_c]
            fecha_formateada = formatear_fecha(informe_actual["created_at"])

            st.divider()
            st.markdown(f"### 📋 Paciente: **{informe_actual['paciente']}**")
            st.caption(f"🕒 Fecha/Hora local: {fecha_formateada}")

            col_txt, col_aud = st.columns(2)
            with col_txt:
                st.subheader("📄 Transcripción Médica")

                with st.form("form_edicion"):
                    texto_editado = st.text_area(
                        "Dictado procesado (Edita texto/apellidos aquí):",
                        informe_actual["transcripcion"],
                        height=220,
                    )

                    guardar = st.form_submit_button(
                        "💾 Guardar cambios y actualizar paciente",
                        type="primary",
                        use_container_width=True,
                    )

                    if guardar:
                        # Extraer el nuevo nombre del texto editado
                        match = re.search(
                            r"Paciente\s+(.*?)(?:,|\s+DNI|\s+del|\s+con|\s+de)",
                            texto_editado,
                            re.IGNORECASE,
                        )
                        nuevo_nombre = (
                            match.group(1).strip().title()
                            if match
                            else informe_actual["paciente"]
                        )

                        # Actualizar transcripción y paciente simultáneamente
                        supabase.table("informes").update({
                            "transcripcion": texto_editado,
                            "paciente": nuevo_nombre,
                        }).eq("id", informe_actual["id"]).execute()

                        st.success(
                            "¡Expediente y nombre actualizados correctamente!"
                        )
                        st.rerun()

                # Botón de Word sincronizado con el estado actual
                word_file = generar_documento_word(
                    informe_actual["paciente"],
                    fecha_formateada,
                    informe_actual["transcripcion"],
                )
                nombre_doc_limpio = sanitizar_nombre_archivo(
                    informe_actual["paciente"]
                )
                st.download_button(
                    label="📥 Descargar Informe en Word (.docx)",
                    data=word_file,
                    file_name=f"Informe_{nombre_doc_limpio}_{informe_actual['id']}.docx",
                    mime=(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ),
                    use_container_width=True,
                )

            with col_aud:
                st.subheader("🔊 Audio Original")
                st.audio(informe_actual["audio_url"])

            st.divider()
            # Portafolio de imágenes médicas
            st.subheader("🖼️ Portafolio de Imágenes Médicas (Rayos X, ECO, TC)")

            imagenes_existentes = informe_actual.get("imagenes_urls") or []
            if isinstance(imagenes_existentes, str):
                try:
                    imagenes_existentes = json.loads(imagenes_existentes)
                except Exception:
                    imagenes_existentes = []

            imagenes_subidas = st.file_uploader(
                "Adjuntar imágenes a esta consulta:",
                type=["png", "jpg", "jpeg", "webp"],
                accept_multiple_files=True,
            )

            if imagenes_subidas and st.button("💾 Guardar Imágenes"):
                nuevas_urls = list(imagenes_existentes)
                for img in imagenes_subidas:
                    img_bytes = img.getvalue()
                    timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                    nombre_img_limpio = sanitizar_nombre_archivo(img.name)
                    img_name = f"img_{informe_actual['id']}_{timestamp_str}_{nombre_img_limpio}"

                    supabase.storage.from_("audios").upload(
                        img_name,
                        img_bytes,
                        file_options={
                            "content-type": img.type,
                            "upsert": "true",
                        },
                    )
                    img_url = supabase.storage.from_("audios").get_public_url(
                        img_name
                    )
                    nuevas_urls.append(img_url)

                supabase.table("informes").update(
                    {"imagenes_urls": json.dumps(nuevas_urls)}
                ).eq("id", informe_actual["id"]).execute()

                st.success("¡Imágenes agregadas al expediente!")
                st.rerun()

            if imagenes_existentes:
                st.markdown("**Imágenes guardadas en esta consulta:**")
                cols = st.columns(3)
                for i, url_img in enumerate(imagenes_existentes):
                    cols[i % 3].image(
                        url_img,
                        caption=f"Imagen {i+1}",
                        use_container_width=True,
                    )
            else:
                st.info("No hay imágenes adjuntas a esta consulta.")

    except Exception as e:
        st.error(f"❌ Error de consulta:\n\n`{e}`")
